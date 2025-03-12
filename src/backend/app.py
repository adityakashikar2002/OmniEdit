from flask import Flask, request, send_file, jsonify
import os
from pdf2docx import Converter
import pdfkit
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
import logging

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# --- pdfkit configuration ---
wkhtmltopdf_path = r"F:\wkhtmltopdf\bin\wkhtmltopdf.exe"  # Correct path
config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
# --- end pdfkit configuration ---

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
def docx_to_html(docx_path):
    """Converts DOCX to HTML preserving formatting"""
    try:
        logging.info(f"docx_path: {docx_path}")
        doc = Document(docx_path)
        logging.info(f"doc: {doc}")
        logging.info(f"doc.paragraphs: {doc.paragraphs}")
        html_content = "<html><body>"
        for para in doc.paragraphs:
            html_content += f"<p>{para.text}</p>"
        html_content += "</body></html>"
        return html_content
    except Exception as e:
        logging.error(f"Error converting DOCX: {str(e)}")
        return f"<p>Error converting DOCX: {str(e)}</p>"

@app.route("/convert_pdf_to_docx", methods=["POST"])
def convert_pdf_to_docx():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    pdf_file = request.files["file"]
    filename = secure_filename(pdf_file.filename)
    pdf_path = os.path.join(UPLOAD_FOLDER, filename)
    docx_path = pdf_path.replace(".pdf", ".docx")

    pdf_file.save(pdf_path)

    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    except Exception as e:
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

    return send_file(docx_path, as_attachment=True)

@app.route("/convert_docx_to_pdf", methods=["POST"])
def convert_docx_to_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    docx_file = request.files["file"]
    filename = secure_filename(docx_file.filename)
    docx_path = os.path.join(UPLOAD_FOLDER, filename)
    pdf_path = docx_path.replace(".docx", ".pdf")
    html_path = docx_path.replace(".docx", ".html")

    docx_file.save(docx_path)

    try:
        # Convert DOCX to HTML
        html_content = docx_to_html(docx_path)

        # Save HTML to a temporary file
        with open(html_path, "w", encoding="utf-8") as html_file:
            html_file.write(html_content)

        # Log HTML content and file size
        logging.info(f"HTML file created: {html_path}")
        logging.info(f"HTML file size: {os.path.getsize(html_path)} bytes")
        logging.info(f"HTML Content:\n{html_content}")

        # Convert HTML to PDF
        logging.info(f"Converting HTML to PDF: {html_path} -> {pdf_path}")
        pdfkit.from_file(html_path, pdf_path, configuration=config)

        # Log PDF file size
        logging.info(f"PDF file created: {pdf_path}")
        logging.info(f"PDF file size: {os.path.getsize(pdf_path)} bytes")

    except FileNotFoundError:
        logging.error("wkhtmltopdf not found")
        return jsonify({"error": "wkhtmltopdf not found"}), 500
    except OSError as e:
        logging.error(f"OS error: {str(e)}")
        return jsonify({"error": f"OS error: {str(e)}"}), 500
    except Exception as e:
        logging.error(f"PDF conversion failed: {str(e)}")
        return jsonify({"error": f"PDF conversion failed: {str(e)}"}), 500

    return send_file(pdf_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True, port=5000)


